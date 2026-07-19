"""
Document Embeddings Module - handle non-PDF documents (Word, CSV, Excel).
"""

import csv
import logging
from pathlib import Path
from typing import Any

from django.utils import timezone
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

logger = logging.getLogger(__name__)


def _read_docx(file_path: Path) -> str:
    """Extract text from a .docx file, including tables."""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    parts = []

    # Paragraphs
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())

    # Tables — convert to pipe-delimited text so LLM can parse them
    for ti, table in enumerate(doc.tables):
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(cells))
        if table_rows:
            parts.append(f"\n[Table {ti + 1}]\n" + "\n".join(table_rows))

    return "\n\n".join(parts)


def _read_doc(file_path: Path) -> str:
    """Extract text from a legacy .doc file using textract if available."""
    try:
        import textract
    except ImportError as exc:
        raise ImportError(
            "Processing .doc files requires the textract package. Install textract to enable .doc ingestion."
        ) from exc

    return textract.process(str(file_path)).decode("utf-8", errors="ignore")


def _read_csv(file_path: Path) -> str:
    """Extract text from a CSV file."""
    rows: list[str] = []
    with open(file_path, newline="", encoding="utf-8", errors="ignore") as handle:
        reader = csv.reader(handle)
        for row in reader:
            row_values = [col.strip() for col in row if col not in (None, "")]
            if row_values:
                rows.append(", ".join(row_values))
    return "\n".join(rows)


def _read_xlsx(file_path: Path) -> str:
    """Extract text from an .xlsx file using openpyxl."""
    try:
        import openpyxl
    except ImportError as exc:
        raise ImportError(
            "Processing .xlsx files requires the openpyxl package. Install openpyxl to enable Excel ingestion."
        ) from exc

    workbook = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            row_values = [str(cell).strip() for cell in row if cell not in (None, "")]
            if row_values:
                lines.append(", ".join(row_values))
    return "\n".join(lines)


def _read_xls(file_path: Path) -> str:
    """Extract text from an .xls file using xlrd."""
    try:
        import xlrd
    except ImportError as exc:
        raise ImportError(
            "Processing .xls files requires the xlrd package. Install xlrd to enable Excel ingestion."
        ) from exc

    workbook = xlrd.open_workbook(file_path)
    lines: list[str] = []
    for sheet in workbook.sheets():
        for row_idx in range(sheet.nrows):
            row_values = [str(cell).strip() for cell in sheet.row_values(row_idx) if str(cell).strip()]
            if row_values:
                lines.append(", ".join(row_values))
    return "\n".join(lines)


def _extract_text(file_path: Path) -> str:
    """Dispatch to the correct extractor based on file extension."""
    ext = file_path.suffix.lower()
    if ext == ".docx":
        return _read_docx(file_path)
    if ext == ".doc":
        return _read_doc(file_path)
    if ext == ".csv":
        return _read_csv(file_path)
    if ext == ".xlsx":
        return _read_xlsx(file_path)
    if ext == ".xls":
        return _read_xls(file_path)
    raise ValueError(f"Unsupported document type: {ext}")


def create_embeddings_for_document(
    file_id: str,
    file_path: str,
    user_id: str = None,
    workspace_id: str = None,
) -> dict[str, Any]:
    """Generate embeddings for non-PDF documents."""
    path = Path(file_path)
    logger.info("🚀 Starting document embeddings creation for file %s", file_id)
    text_content = _extract_text(path)

    if not text_content or not text_content.strip():
        raise ValueError(f"No readable text found in document {file_id}")

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100,
    )

    base_metadata = {
        "pdf_id": file_id,  # Retain key for compatibility with existing chat flows
        "user_id": user_id,
        "workspace_id": workspace_id,
        "type": path.suffix.lstrip(".") or "document",
        "status": "active",
        "created_at": timezone.now().isoformat(),
        "privacy": "private",
        "source": path.name,
    }

    documents = text_splitter.split_documents([Document(page_content=text_content, metadata=base_metadata)])

    for idx, doc in enumerate(documents):
        doc.metadata.update(
            {
                "page": idx + 1,
                "text": doc.page_content,
            }
        )

    # Store where agent retrieval reads (ai_embedding_chunks), not LangChain
    # PGVector's own tables — same store-split fix as PDFs (2026-07-15). Without
    # this, a chatted .docx/.csv would also return "No content found".
    from components.knowledge.infrastructure.adapters.pgvector_document_indexer import (
        index_documents,
    )

    index_documents(documents)

    logger.info("✅ Stored %s document chunks for file %s", len(documents), file_id)

    return {
        "success": True,
        "file_id": file_id,
        "file_path": file_path,
        "chunks_created": len(documents),
        "embeddings_generated": len(documents),
        "chunk_size": 500,
        "chunk_overlap": 100,
        "file_type": base_metadata["type"],
    }
